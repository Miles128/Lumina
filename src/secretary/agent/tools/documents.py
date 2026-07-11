"""Read Excel / PDF / Word documents as text for the agent."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from secretary.agent.text_utils import truncate_chars
from secretary.agent.tools.base import Tool, _resolve_path

_MAX_OUTPUT_CHARS = 24_000
_SUPPORTED = {".xlsx", ".xlsm", ".pdf", ".docx"}


def _truncate(text: str, limit: int = _MAX_OUTPUT_CHARS) -> str:
    cleaned = text.strip() or "(empty document)"
    return truncate_chars(cleaned, limit)


def _read_xlsx(path: Path, *, sheet: str | None, max_rows: int) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl is not installed (required for Excel)"

    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        if sheet:
            if sheet not in wb.sheetnames:
                return f"Error: sheet not found: {sheet}. Available: {', '.join(wb.sheetnames)}"
            sheets = [wb[sheet]]
        else:
            sheets = list(wb.worksheets)

        blocks: list[str] = []
        for ws in sheets:
            lines = [f"## Sheet: {ws.title}"]
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows:
                    lines.append(f"...[truncated at {max_rows} rows]")
                    break
                cells = [("" if cell is None else str(cell)) for cell in row]
                if any(cells):
                    lines.append("\t".join(cells))
                    row_count += 1
            blocks.append("\n".join(lines))
        return _truncate("\n\n".join(blocks))
    finally:
        wb.close()


def _read_pdf(path: Path, *, page_start: int, page_end: int | None) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "Error: pypdf is not installed (required for PDF)"

    reader = PdfReader(str(path))
    total = len(reader.pages)
    if total == 0:
        return "(empty PDF)"
    start = max(1, page_start) - 1
    end = total if page_end is None else min(total, max(page_start, page_end))
    if start >= total:
        return f"Error: page_start {page_start} exceeds page count {total}"

    parts: list[str] = [f"## PDF ({total} pages), extracting {start + 1}–{end}"]
    for idx in range(start, end):
        text = (reader.pages[idx].extract_text() or "").strip()
        parts.append(f"### Page {idx + 1}\n{text or '(no extractable text)'}")
    return _truncate("\n\n".join(parts))


def _read_docx(path: Path, *, max_paragraphs: int) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx is not installed (required for Word)"

    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    truncated = False
    if len(paras) > max_paragraphs:
        paras = paras[:max_paragraphs]
        truncated = True

    # Include simple tables
    table_blocks: list[str] = []
    for ti, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("\t".join(cells))
        if rows:
            table_blocks.append(f"## Table {ti}\n" + "\n".join(rows))

    parts = ["## Body"] + paras
    if truncated:
        parts.append(f"...[truncated at {max_paragraphs} paragraphs]")
    if table_blocks:
        parts.extend([""] + table_blocks)
    return _truncate("\n".join(parts))


class ReadDocumentTool(Tool):
    name = "read_document"
    description = (
        "Extract text/tables from Excel (.xlsx), PDF (.pdf), or Word (.docx). "
        "Use this instead of file_read for binary office documents."
    )
    needs_confirmation = False
    risk_level = "low"

    def _parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Path to .xlsx / .pdf / .docx"},
                "sheet": {
                    "type": "string",
                    "description": "Excel sheet name (default: all sheets)",
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Max rows per Excel sheet (default 200)",
                },
                "page_start": {
                    "type": "integer",
                    "description": "PDF first page (1-based, default 1)",
                },
                "page_end": {
                    "type": "integer",
                    "description": "PDF last page inclusive (default: all)",
                },
                "max_paragraphs": {
                    "type": "integer",
                    "description": "Max Word paragraphs (default 400)",
                },
            },
            "required": ["path"],
        }

    def describe_action(self, arguments: dict[str, Any], working_dir: Path) -> str:
        path = arguments.get("path", "")
        return f"📄 读取文档: `{path}`"

    def execute(self, arguments: dict[str, Any], working_dir: Path) -> str:
        raw = str(arguments.get("path", "")).strip()
        if not raw:
            return "Error: path is required"
        path = _resolve_path(raw, working_dir)
        if not path.exists():
            return f"Error: path not found: {path}"
        if not path.is_file():
            return f"Error: not a file: {path}"

        suffix = path.suffix.lower()
        if suffix not in _SUPPORTED:
            return (
                f"Error: unsupported format '{suffix}'. "
                f"Supported: {', '.join(sorted(_SUPPORTED))}"
            )

        try:
            if suffix in {".xlsx", ".xlsm"}:
                sheet = str(arguments.get("sheet") or "").strip() or None
                max_rows = min(int(arguments.get("max_rows", 200) or 200), 2000)
                return _read_xlsx(path, sheet=sheet, max_rows=max_rows)
            if suffix == ".pdf":
                page_start = int(arguments.get("page_start", 1) or 1)
                page_end_raw = arguments.get("page_end")
                page_end = int(page_end_raw) if page_end_raw is not None else None
                return _read_pdf(path, page_start=page_start, page_end=page_end)
            max_paragraphs = min(int(arguments.get("max_paragraphs", 400) or 400), 2000)
            return _read_docx(path, max_paragraphs=max_paragraphs)
        except Exception as exc:
            return f"Error: failed to read {path.name}: {exc}"
